from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from pydantic import BaseModel
from datetime import datetime, timezone
import uuid, os, shutil, io

from app.db.session import get_db
from app.db.models import PocReport, PocEvidence, PocRetesting, PocRetestEvidence, VulnReport
from app.api.deps import get_current_user
from app.db.models import User

router = APIRouter()
UPLOAD_DIR = "/app/uploads/poc_evidence"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ── Schemas ───────────────────────────────────────────────────
class PocCreate(BaseModel):
    vuln_report_id: str
    status: str = "BELUM DIPERBAIKI"
    description: Optional[str] = None
    poc_steps: Optional[str] = None
    reference: Optional[str] = None
    recommendation: Optional[str] = None

class PocUpdate(BaseModel):
    status: Optional[str] = None
    description: Optional[str] = None
    poc_steps: Optional[str] = None
    reference: Optional[str] = None
    recommendation: Optional[str] = None

class RetestCreate(BaseModel):
    poc_id: str
    retest_date: Optional[str] = None
    result: Optional[str] = None
    status: str = "BELUM DIPERBAIKI"

# ── Helper ────────────────────────────────────────────────────
def _poc_dict(p: PocReport) -> dict:
    return {
        "id": p.id,
        "vuln_report_id": p.vuln_report_id,
        "vuln_id": p.vuln_id,
        "vuln_name": p.vuln_name,
        "cvss_vector": p.cvss_vector,
        "cvss_score": p.cvss_score,
        "severity": p.severity,
        "status": p.status,
        "description": p.description,
        "poc_steps": p.poc_steps,
        "reference": p.reference,
        "recommendation": p.recommendation,
        "created_at": p.created_at.isoformat() if p.created_at else None,
        "updated_at": p.updated_at.isoformat() if p.updated_at else None,
        "evidences": [_ev_dict(e) for e in (p.evidences or [])],
        "retestings": [_rt_dict(r) for r in (p.retestings or [])],
    }

def _ev_dict(e: PocEvidence) -> dict:
    return {
        "id": e.id, "poc_id": e.poc_id,
        "order_no": e.order_no, "label": e.label,
        "caption": e.caption,
        "file_path": e.file_path,
        "file_name": e.file_name,
        "url": f"/uploads/poc_evidence/{os.path.basename(e.file_path)}" if e.file_path else None,
    }

def _rt_dict(r: PocRetesting) -> dict:
    return {
        "id": r.id, "poc_id": r.poc_id,
        "retest_date": r.retest_date.strftime("%Y-%m-%d") if r.retest_date else None,
        "result": r.result, "status": r.status,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "evidences": [_rte_dict(e) for e in (r.evidences or [])],
    }

def _rte_dict(e: PocRetestEvidence) -> dict:
    return {
        "id": e.id, "retest_id": e.retest_id,
        "order_no": e.order_no, "label": e.label,
        "caption": e.caption, "file_name": e.file_name,
        "url": f"/uploads/poc_evidence/{os.path.basename(e.file_path)}" if e.file_path else None,
    }

# ── POC CRUD ──────────────────────────────────────────────────
@router.get("/poc")
async def list_pocs(
    vuln_report_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    q = select(PocReport).order_by(PocReport.created_at.desc())
    if vuln_report_id:
        q = q.where(PocReport.vuln_report_id == vuln_report_id)
    result = await db.execute(q)
    pocs = result.scalars().all()
    out = []
    for p in pocs:
        ev_q = await db.execute(select(PocEvidence).where(PocEvidence.poc_id == p.id).order_by(PocEvidence.order_no))
        evidences = ev_q.scalars().all()
        rt_q = await db.execute(select(PocRetesting).where(PocRetesting.poc_id == p.id).order_by(PocRetesting.created_at))
        rts = rt_q.scalars().all()
        rt_list = []
        for rt in rts:
            rte_q = await db.execute(select(PocRetestEvidence).where(PocRetestEvidence.retest_id == rt.id))
            rte_list = rte_q.scalars().all()
            rt_list.append({
                "id": rt.id, "poc_id": rt.poc_id,
                "retest_date": rt.retest_date.strftime("%Y-%m-%d") if rt.retest_date else None,
                "result": rt.result, "status": rt.status,
                "created_at": rt.created_at.isoformat() if rt.created_at else None,
                "evidences": [_rte_dict(e) for e in rte_list],
            })
        out.append({
            "id": p.id,
            "vuln_report_id": p.vuln_report_id,
            "vuln_id": p.vuln_id,
            "vuln_name": p.vuln_name,
            "cvss_vector": p.cvss_vector,
            "cvss_score": p.cvss_score,
            "severity": p.severity,
            "status": p.status,
            "description": p.description,
            "poc_steps": p.poc_steps,
            "reference": p.reference,
            "recommendation": p.recommendation,
            "created_at": p.created_at.isoformat() if p.created_at else None,
            "updated_at": p.updated_at.isoformat() if p.updated_at else None,
            "evidences": [_ev_dict(e) for e in evidences],
            "retestings": rt_list,
        })
    return out

@router.post("/poc")
async def create_poc(
    data: PocCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Auto-fill from vuln_report
    vr = await db.execute(select(VulnReport).where(VulnReport.id == data.vuln_report_id))
    vuln = vr.scalar_one_or_none()
    if not vuln:
        raise HTTPException(404, "Vuln report not found")

    p = PocReport(
        id=str(uuid.uuid4()),
        vuln_report_id=data.vuln_report_id,
        vuln_id=vuln.vuln_id,
        vuln_name=vuln.vuln_name,
        cvss_vector=vuln.cvss_vector,
        cvss_score=vuln.cvss_score,
        severity=vuln.severity,
        status=data.status,
        description=data.description,
        poc_steps=data.poc_steps,
        reference=vuln.referensi,  # auto-fill
        recommendation=data.recommendation,
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc),
    )
    db.add(p)
    await db.commit()
    # Return dict directly without touching relationships
    return {
        "id": p.id,
        "vuln_report_id": p.vuln_report_id,
        "vuln_id": p.vuln_id,
        "vuln_name": p.vuln_name,
        "cvss_vector": p.cvss_vector,
        "cvss_score": p.cvss_score,
        "severity": p.severity,
        "status": p.status,
        "description": p.description,
        "poc_steps": p.poc_steps,
        "reference": p.reference,
        "recommendation": p.recommendation,
        "created_at": p.created_at.isoformat() if p.created_at else None,
        "updated_at": p.updated_at.isoformat() if p.updated_at else None,
        "evidences": [],
        "retestings": [],
    }

@router.get("/poc/{poc_id}")
async def get_poc(
    poc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    p = await db.execute(select(PocReport).where(PocReport.id == poc_id))
    poc = p.scalar_one_or_none()
    if not poc: raise HTTPException(404, "Not found")

    ev_q = await db.execute(select(PocEvidence).where(PocEvidence.poc_id == poc_id).order_by(PocEvidence.order_no))
    evidences = ev_q.scalars().all()

    rt_q = await db.execute(select(PocRetesting).where(PocRetesting.poc_id == poc_id).order_by(PocRetesting.created_at))
    rts = rt_q.scalars().all()

    rt_list = []
    for rt in rts:
        rte_q = await db.execute(select(PocRetestEvidence).where(PocRetestEvidence.retest_id == rt.id))
        rte_list = rte_q.scalars().all()
        rt_list.append({
            "id": rt.id, "poc_id": rt.poc_id,
            "retest_date": rt.retest_date.strftime("%Y-%m-%d") if rt.retest_date else None,
            "result": rt.result, "status": rt.status,
            "created_at": rt.created_at.isoformat() if rt.created_at else None,
            "evidences": [_rte_dict(e) for e in rte_list],
        })

    return {
        "id": poc.id,
        "vuln_report_id": poc.vuln_report_id,
        "vuln_id": poc.vuln_id,
        "vuln_name": poc.vuln_name,
        "cvss_vector": poc.cvss_vector,
        "cvss_score": poc.cvss_score,
        "severity": poc.severity,
        "status": poc.status,
        "description": poc.description,
        "poc_steps": poc.poc_steps,
        "reference": poc.reference,
        "recommendation": poc.recommendation,
        "created_at": poc.created_at.isoformat() if poc.created_at else None,
        "updated_at": poc.updated_at.isoformat() if poc.updated_at else None,
        "evidences": [_ev_dict(e) for e in evidences],
        "retestings": rt_list,
    }

@router.put("/poc/{poc_id}")
async def update_poc(
    poc_id: str,
    data: PocUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    result = await db.execute(select(PocReport).where(PocReport.id == poc_id))
    poc = result.scalar_one_or_none()
    if not poc: raise HTTPException(404, "Not found")
    for k, v in data.model_dump(exclude_none=True).items():
        setattr(poc, k, v)
    poc.updated_at = datetime.now(timezone.utc)
    await db.commit()
    return {"ok": True}

@router.delete("/poc/{poc_id}")
async def delete_poc(
    poc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Delete evidence files
    evs = await db.execute(select(PocEvidence).where(PocEvidence.poc_id == poc_id))
    for ev in evs.scalars():
        if ev.file_path and os.path.exists(ev.file_path):
            os.remove(ev.file_path)
    await db.execute(delete(PocEvidence).where(PocEvidence.poc_id == poc_id))
    await db.execute(delete(PocReport).where(PocReport.id == poc_id))
    await db.commit()
    return {"ok": True}

# ── Evidence Upload ───────────────────────────────────────────
@router.post("/poc/{poc_id}/evidence")
async def upload_evidence(
    poc_id: str,
    file: UploadFile = File(...),
    caption: str = Form(""),
    label: str = Form("Evidence-01"),
    order_no: int = Form(1),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # Validate image
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in [".png", ".jpg", ".jpeg"]:
        raise HTTPException(400, "Only PNG, JPG, JPEG allowed")

    fname = f"{uuid.uuid4()}{ext}"
    fpath = os.path.join(UPLOAD_DIR, fname)
    with open(fpath, "wb") as f_out:
        shutil.copyfileobj(file.file, f_out)

    ev = PocEvidence(
        id=str(uuid.uuid4()),
        poc_id=poc_id,
        order_no=order_no,
        label=label,
        caption=caption,
        file_path=fpath,
        file_name=file.filename,
        created_at=datetime.now(timezone.utc),
    )
    db.add(ev)
    await db.commit()
    return _ev_dict(ev)

@router.delete("/poc/evidence/{evidence_id}")
async def delete_evidence(
    evidence_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ev_q = await db.execute(select(PocEvidence).where(PocEvidence.id == evidence_id))
    ev = ev_q.scalar_one_or_none()
    if ev and ev.file_path and os.path.exists(ev.file_path):
        os.remove(ev.file_path)
    await db.execute(delete(PocEvidence).where(PocEvidence.id == evidence_id))
    await db.commit()
    return {"ok": True}

# ── Retesting ─────────────────────────────────────────────────
@router.post("/poc/{poc_id}/retest")
async def create_retest(
    poc_id: str,
    data: RetestCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from app.api.routes.vuln_mgmt import parse_date
    rt = PocRetesting(
        id=str(uuid.uuid4()),
        poc_id=poc_id,
        retest_date=parse_date(data.retest_date),
        result=data.result,
        status=data.status,
        created_at=datetime.now(timezone.utc),
    )
    db.add(rt)
    await db.commit()
    return {
        "id": rt.id,
        "poc_id": rt.poc_id,
        "retest_date": rt.retest_date.strftime("%Y-%m-%d") if rt.retest_date else None,
        "result": rt.result,
        "status": rt.status,
        "created_at": rt.created_at.isoformat() if rt.created_at else None,
        "evidences": [],
    }

@router.put("/poc/retest/{retest_id}")
async def update_retest(
    retest_id: str,
    result: Optional[str] = None,
    status: Optional[str] = None,
    retest_date: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from app.api.routes.vuln_mgmt import parse_date
    rt_q = await db.execute(select(PocRetesting).where(PocRetesting.id == retest_id))
    rt = rt_q.scalar_one_or_none()
    if not rt: raise HTTPException(404, "Not found")
    if result is not None: rt.result = result
    if status is not None: rt.status = status
    if retest_date is not None: rt.retest_date = parse_date(retest_date)
    await db.commit()
    return {"ok": True}

@router.delete("/poc/retest/{retest_id}")
async def delete_retest(
    retest_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    await db.execute(delete(PocRetestEvidence).where(PocRetestEvidence.retest_id == retest_id))
    await db.execute(delete(PocRetesting).where(PocRetesting.id == retest_id))
    await db.commit()
    return {"ok": True}

@router.post("/poc/retest/{retest_id}/evidence")
async def upload_retest_evidence(
    retest_id: str,
    file: UploadFile = File(...),
    caption: str = Form(""),
    label: str = Form("Evidence-01"),
    order_no: int = Form(1),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in [".png", ".jpg", ".jpeg"]:
        raise HTTPException(400, "Only PNG, JPG, JPEG allowed")
    fname = f"{uuid.uuid4()}{ext}"
    fpath = os.path.join(UPLOAD_DIR, fname)
    with open(fpath, "wb") as f_out:
        shutil.copyfileobj(file.file, f_out)
    ev = PocRetestEvidence(
        id=str(uuid.uuid4()),
        retest_id=retest_id,
        order_no=order_no,
        label=label,
        caption=caption,
        file_path=fpath,
        file_name=file.filename,
        created_at=datetime.now(timezone.utc),
    )
    db.add(ev)
    await db.commit()
    return _rte_dict(ev)

# ── Export PDF ────────────────────────────────────────────────
@router.get("/poc/{poc_id}/export-pdf")
async def export_poc_pdf(
    poc_id: str,
    token: str = "",
    db: AsyncSession = Depends(get_db),
):
    from app.core.security import decode_token
    try:
        payload = decode_token(token)
        user_id = payload.get("sub")
        if not user_id: raise Exception("Invalid token")
    except:
        from fastapi import HTTPException as _HTTPException
        raise _HTTPException(401, "Not authenticated")
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    # Get POC
    p_q = await db.execute(select(PocReport).where(PocReport.id == poc_id))
    poc = p_q.scalar_one_or_none()
    if not poc: raise HTTPException(404, "Not found")
    ev_q = await db.execute(select(PocEvidence).where(PocEvidence.poc_id == poc_id).order_by(PocEvidence.order_no))
    evidences_pdf = list(ev_q.scalars().all())
    rt_q = await db.execute(select(PocRetesting).where(PocRetesting.poc_id == poc_id).order_by(PocRetesting.created_at))
    rts_pdf = list(rt_q.scalars().all())
    rts_with_ev = []
    for rt in rts_pdf:
        rte_q = await db.execute(select(PocRetestEvidence).where(PocRetestEvidence.retest_id == rt.id))
        rte_list = list(rte_q.scalars().all())
        rts_with_ev.append((rt, rte_list))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    sev_colors = {
        'critical': colors.HexColor('#ff5f5f'),
        'high': colors.HexColor('#ff9f43'),
        'medium': colors.HexColor('#ffd43b'),
        'low': colors.HexColor('#a9e34b'),
        'informational': colors.HexColor('#74c7ec'),
    }
    sev_color = sev_colors.get((poc.severity or 'medium').lower(), colors.HexColor('#74c7ec'))
    # Use local vars instead of relationship attrs
    poc_evidences = evidences_pdf
    poc_retestings = rts_with_ev
    status_color = colors.HexColor('#ff5f5f') if poc.status == 'BELUM DIPERBAIKI' else colors.HexColor('#a9e34b')

    title_style = ParagraphStyle('title', fontSize=16, fontName='Helvetica-Bold', spaceAfter=6, textColor=colors.HexColor('#1e1e2e'))
    heading_style = ParagraphStyle('heading', fontSize=11, fontName='Helvetica-Bold', spaceAfter=4, spaceBefore=12, textColor=colors.HexColor('#313244'))
    body_style = ParagraphStyle('body', fontSize=9, fontName='Helvetica', spaceAfter=4, leading=14)
    label_style = ParagraphStyle('label', fontSize=9, fontName='Helvetica-Bold', textColor=colors.HexColor('#6c7086'))

    story = []

    # Title
    story.append(Paragraph("Proof of Concept (POC)", title_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#6366f1')))
    story.append(Spacer(1, 0.3*cm))

    # Info table
    info_data = [
        [Paragraph("Vulnerability ID", label_style), Paragraph(poc.vuln_id or '-', body_style)],
        [Paragraph("Vulnerability Name", label_style), Paragraph(poc.vuln_name or '-', body_style)],
        [Paragraph("CVSS Vector", label_style), Paragraph(poc.cvss_vector or '-', body_style)],
        [Paragraph("CVSS Score", label_style), Paragraph(str(poc.cvss_score or '-'), body_style)],
        [Paragraph("Severity", label_style), Paragraph((poc.severity or '-').upper(), ParagraphStyle('sev', fontSize=9, fontName='Helvetica-Bold', textColor=sev_color))],
        [Paragraph("Status", label_style), Paragraph(poc.status or '-', ParagraphStyle('st', fontSize=9, fontName='Helvetica-Bold', textColor=status_color))],
    ]
    info_table = Table(info_data, colWidths=[4*cm, 13*cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f0f0f0')),
        ('ROWBACKGROUNDS', (0,0), (-1,-1), [colors.white, colors.HexColor('#fafafa')]),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e0e0e0')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(info_table)

    # Description
    if poc.description:
        story.append(Paragraph("Description", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 0.2*cm))
        for line in poc.description.split('\n'):
            story.append(Paragraph(line or '&nbsp;', body_style))

    # POC Steps
    if poc.poc_steps:
        story.append(Paragraph("Proof of Concept", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 0.2*cm))
        for line in poc.poc_steps.split('\n'):
            story.append(Paragraph(line or '&nbsp;', body_style))

    # Evidences
    if poc_evidences:
        story.append(Paragraph("Evidence", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 0.2*cm))
        for ev in poc_evidences:
            story.append(Paragraph(f"<b>{ev.label}</b>", body_style))
            if ev.caption:
                story.append(Paragraph(ev.caption, body_style))
            if ev.file_path and os.path.exists(ev.file_path):
                try:
                    img = RLImage(ev.file_path, width=15*cm, height=8*cm, kind='proportional')
                    story.append(img)
                except: pass
            story.append(Spacer(1, 0.3*cm))

    # Reference
    if poc.reference:
        story.append(Paragraph("Reference", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 0.2*cm))
        for line in poc.reference.split('\n'):
            story.append(Paragraph(line or '&nbsp;', body_style))

    # Recommendation
    if poc.recommendation:
        story.append(Paragraph("Recommendation", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
        story.append(Spacer(1, 0.2*cm))
        for line in poc.recommendation.split('\n'):
            story.append(Paragraph(line or '&nbsp;', body_style))

    # Retesting
    for i, (rt, rt_evidences) in enumerate(poc_retestings):
        rt_status_color = colors.HexColor('#ff5f5f') if rt.status == 'BELUM DIPERBAIKI' else colors.HexColor('#a9e34b')
        story.append(Paragraph(f"Hasil Retesting #{i+1}", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#6366f1')))
        rt_data = [
            [Paragraph("Tanggal", label_style), Paragraph(rt.retest_date or '-', body_style)],
            [Paragraph("Status", label_style), Paragraph(rt.status or '-', ParagraphStyle('rts', fontSize=9, fontName='Helvetica-Bold', textColor=rt_status_color))],
            [Paragraph("Hasil", label_style), Paragraph(rt.result or '-', body_style)],
        ]
        rt_table = Table(rt_data, colWidths=[4*cm, 13*cm])
        rt_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f0f0f0')),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e0e0e0')),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(rt_table)
        for ev in rt_evidences:
            story.append(Paragraph(f"<b>{ev.label or 'Evidence'}</b>", body_style))
            if ev.caption: story.append(Paragraph(ev.caption, body_style))
            if ev.file_path and os.path.exists(ev.file_path):
                try:
                    img = RLImage(ev.file_path, width=15*cm, height=8*cm, kind='proportional')
                    story.append(img)
                except: pass
            story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    buf.seek(0)
    filename = f"POC_{poc.vuln_id or poc_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"})


# ── Export Word ───────────────────────────────────────────────
@router.get("/poc/{poc_id}/export-word")
async def export_poc_word(
    poc_id: str,
    token: str = "",
    db: AsyncSession = Depends(get_db),
):
    from app.core.security import decode_token
    try:
        payload = decode_token(token)
        user_id = payload.get("sub")
        if not user_id: raise Exception("Invalid token")
    except:
        from fastapi import HTTPException as _HTTPException
        raise _HTTPException(401, "Not authenticated")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Get POC
    p_q = await db.execute(select(PocReport).where(PocReport.id == poc_id))
    poc = p_q.scalar_one_or_none()
    if not poc: raise HTTPException(404, "Not found")
    ev_q = await db.execute(select(PocEvidence).where(PocEvidence.poc_id == poc_id).order_by(PocEvidence.order_no))
    evidences = ev_q.scalars().all()
    rt_q = await db.execute(select(PocRetesting).where(PocRetesting.poc_id == poc_id).order_by(PocRetesting.created_at))
    rts = rt_q.scalars().all()
    for rt in rts:
        rte_q = await db.execute(select(PocRetestEvidence).where(PocRetestEvidence.retest_id == rt.id))
        rt.evidences = rte_q.scalars().all()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    sev_colors = {
        'critical': RGBColor(0xFF, 0x5F, 0x5F),
        'high':     RGBColor(0xFF, 0x9F, 0x43),
        'medium':   RGBColor(0xFF, 0xD4, 0x3B),
        'low':      RGBColor(0xA9, 0xE3, 0x4B),
        'informational': RGBColor(0x74, 0xC7, 0xEC),
    }
    sev_color = sev_colors.get((poc.severity or 'medium').lower(), RGBColor(0x74, 0xC7, 0xEC))
    status_color = RGBColor(0xA9, 0xE3, 0x4B) if poc.status == 'DIPERBAIKI' else RGBColor(0xFF, 0x5F, 0x5F)

    def add_heading(text, level=1, color=RGBColor(0x31, 0x32, 0x44)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def add_field(label, value, val_color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(0x6C, 0x70, 0x86)
        run_val = p.add_run(str(value or '-'))
        run_val.font.size = Pt(10)
        if val_color:
            run_val.font.color.rgb = val_color
            run_val.bold = True
        return p

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
        # Add border bottom
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6366F1')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Title
    title = doc.add_heading('Proof of Concept (POC)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
        run.font.size = Pt(18)

    doc.add_paragraph()

    # Info fields
    add_field('Vulnerability ID', poc.vuln_id)
    add_field('Vulnerability Name', poc.vuln_name)
    add_field('CVSS Vector', poc.cvss_vector)
    add_field('CVSS Score', poc.cvss_score)
    add_field('Severity', (poc.severity or '').upper(), sev_color)
    add_field('Status', poc.status, status_color)

    # Description
    if poc.description:
        add_section_heading('Description')
        p = doc.add_paragraph(poc.description)
        p.style.font.size = Pt(10)

    # POC Steps
    if poc.poc_steps:
        add_section_heading('Proof of Concept')
        p = doc.add_paragraph(poc.poc_steps)
        p.style.font.size = Pt(10)

    # Evidences
    if evidences:
        add_section_heading('Evidence')
        for ev in evidences:
            p = doc.add_paragraph()
            run = p.add_run(ev.label or 'Evidence')
            run.bold = True
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)
            if ev.caption:
                doc.add_paragraph(ev.caption).style.font.size = Pt(9)
            if ev.file_path and os.path.exists(ev.file_path):
                try:
                    doc.add_picture(ev.file_path, width=Inches(5.5))
                except: pass
            doc.add_paragraph()

    # Reference
    if poc.reference:
        add_section_heading('Reference')
        doc.add_paragraph(poc.reference)

    # Recommendation
    if poc.recommendation:
        add_section_heading('Recommendation')
        doc.add_paragraph(poc.recommendation)

    # Retestings
    for i, rt in enumerate(rts):
        rt_status_color = RGBColor(0xA9, 0xE3, 0x4B) if rt.status == 'DIPERBAIKI' else RGBColor(0xFF, 0x5F, 0x5F)
        add_section_heading(f'Hasil Retesting #{i+1}')
        add_field('Tanggal', rt.retest_date.strftime('%Y-%m-%d') if rt.retest_date else '-')
        add_field('Status', rt.status, rt_status_color)
        if rt.result:
            doc.add_paragraph(rt.result)
        for ev in rt_evidences:
            p = doc.add_paragraph()
            run = p.add_run(ev.label or 'Evidence')
            run.bold = True
            if ev.file_path and os.path.exists(ev.file_path):
                try:
                    doc.add_picture(ev.file_path, width=Inches(5.5))
                except: pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"POC_{poc.vuln_id or poc_id}_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"})
